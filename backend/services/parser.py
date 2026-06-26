"""
Document Parser Service
Extracts text from PDF, DOCX, and TXT files with page tracking.
"""

from pathlib import Path
from typing import List, Dict


def extract_text(filepath: str) -> List[Dict]:
    """
    Extract text from a document file.
    
    Returns:
        List of dicts: [{"text": str, "page": int}, ...]
    """
    ext = Path(filepath).suffix.lower()
    
    if ext == ".pdf":
        return _extract_pdf(filepath)
    elif ext == ".docx":
        return _extract_docx(filepath)
    elif ext == ".txt":
        return _extract_txt(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(filepath: str) -> List[Dict]:
    """Extract text from PDF using PyMuPDF."""
    import fitz  # PyMuPDF
    
    pages = []
    doc = fitz.open(filepath)
    
    for page_num, page in enumerate(doc, start=1):
        text = page.get_text("text")
        if text.strip():  # Skip empty pages
            pages.append({
                "text": text,
                "page": page_num
            })
    
    doc.close()
    return pages


def _extract_docx(filepath: str) -> List[Dict]:
    """Extract text from DOCX using python-docx."""
    from docx import Document
    
    doc = Document(filepath)
    pages = []
    current_text = []
    page_num = 1
    para_count = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            current_text.append(text)
            para_count += 1
            
            # Simulate page breaks every ~50 paragraphs
            if para_count % 50 == 0:
                if current_text:
                    pages.append({
                        "text": "\n".join(current_text),
                        "page": page_num
                    })
                    current_text = []
                    page_num += 1
    
    # Add remaining text
    if current_text:
        pages.append({
            "text": "\n".join(current_text),
            "page": page_num
        })
    
    # Handle tables
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
            if row_text:
                table_text.append(row_text)
        if table_text:
            pages.append({
                "text": "\n".join(table_text),
                "page": page_num
            })
    
    return pages if pages else [{"text": "No content found", "page": 1}]


def _extract_txt(filepath: str) -> List[Dict]:
    """Extract text from TXT file."""
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    
    if not content.strip():
        return []
    
    # Split into pseudo-pages of ~3000 chars
    chunk_size = 3000
    pages = []
    
    for i, start in enumerate(range(0, len(content), chunk_size), start=1):
        text = content[start:start + chunk_size]
        if text.strip():
            pages.append({
                "text": text,
                "page": i
            })
    
    return pages
